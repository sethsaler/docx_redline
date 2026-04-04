from __future__ import annotations

import datetime as _dt

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from docx_redline.doc_walk import list_redline_paragraphs
from docx_redline.differ import Change, ChangeType, DiffSegment, ParagraphDiff, RunInfo

RED = RGBColor(0xFF, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
BLACK = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
HEADER_BG = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_FONT = "Aptos"

TRACK_AUTHOR = "DOCX Redline"


def _enable_track_revisions(doc: Document) -> None:
    settings_el = doc.settings._element
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for el in settings_el.findall(f"{{{ns}}}trackRevisions"):
        settings_el.remove(el)
    tr = OxmlElement("w:trackRevisions")
    settings_el.append(tr)


def _revision_date_iso() -> str:
    return _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _next_revision_id(counter: list[int]) -> str:
    counter[0] += 1
    return str(counter[0])


def _revision_attrs(counter: list[int]) -> dict[str, str]:
    return {
        qn("w:id"): _next_revision_id(counter),
        qn("w:author"): TRACK_AUTHOR,
        qn("w:date"): _revision_date_iso(),
    }


def _deepcopy_rpr(rpr) -> OxmlElement | None:
    if rpr is None:
        return None
    import copy

    return copy.deepcopy(rpr)


def _first_run_rpr_clone(paragraph) -> OxmlElement | None:
    p_el = paragraph._element
    r = p_el.find(qn("w:r"))
    if r is None:
        return None
    return _deepcopy_rpr(r.find(qn("w:rPr")))


def _rpr_from_run_info(run: RunInfo) -> OxmlElement:
    rPr = OxmlElement("w:rPr")
    if run.bold:
        rPr.append(OxmlElement("w:b"))
    if run.italic:
        rPr.append(OxmlElement("w:i"))
    if run.underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    if run.font_name:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), run.font_name)
        fonts.set(qn("w:hAnsi"), run.font_name)
        rPr.append(fonts)
    if run.font_size_pt is not None:
        half_pts = str(int(round(run.font_size_pt * 2)))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), half_pts)
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), half_pts)
        rPr.append(szCs)
    if run.font_color_rgb:
        r, g, b = run.font_color_rgb
        color = OxmlElement("w:color")
        color.set(qn("w:val"), f"{r:02X}{g:02X}{b:02X}")
        rPr.append(color)
    return rPr


def _append_t_to_run(r_elem: OxmlElement, text: str, *, del_text: bool = False) -> None:
    tag = "w:delText" if del_text else "w:t"
    t = OxmlElement(tag)
    t.text = text
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        t.set(qn("xml:space"), "preserve")
    r_elem.append(t)


def _append_ins_run(
    p_elem,
    text: str,
    rpr: OxmlElement | None,
    counter: list[int],
) -> None:
    ins = OxmlElement("w:ins")
    for k, v in _revision_attrs(counter).items():
        ins.set(k, v)
    r = OxmlElement("w:r")
    if rpr is not None and len(rpr):
        r.append(rpr)
    _append_t_to_run(r, text)
    ins.append(r)
    p_elem.append(ins)


def _append_del_run(
    p_elem,
    text: str,
    rpr: OxmlElement | None,
    counter: list[int],
) -> None:
    del_el = OxmlElement("w:del")
    for k, v in _revision_attrs(counter).items():
        del_el.set(k, v)
    r = OxmlElement("w:r")
    if rpr is not None and len(rpr):
        r.append(rpr)
    _append_t_to_run(r, text, del_text=True)
    del_el.append(r)
    p_elem.append(del_el)


def _append_plain_run(p_elem, text: str, rpr: OxmlElement | None) -> None:
    r = OxmlElement("w:r")
    if rpr is not None and len(rpr):
        r.append(rpr)
    _append_t_to_run(r, text)
    p_elem.append(r)


def _add_formatted_run(
    paragraph,
    text: str,
    bold: bool | None = None,
    italic: bool | None = None,
    font_name: str | None = None,
    font_size_pt: float | None = None,
    color: RGBColor | None = None,
    strike: bool = False,
    underline: bool = False,
    highlight: WD_COLOR | None = None,
):
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if font_name:
        run.font.name = font_name
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    if color:
        run.font.color.rgb = color
    if strike:
        run.font.strike = True
    if underline:
        run.font.underline = True
    if highlight:
        run.font.highlight_color = highlight
    return run


def _set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_table_fixed_layout(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def _set_cell_margins(cell, top=0, bottom=0, left=40, right=40):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_mar = tcPr.find(qn("w:tcMar"))
    if existing_mar is not None:
        tcPr.remove(existing_mar)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:w"), str(val))
        elem.set(qn("w:type"), "dxa")
        tcMar.append(elem)
    tcPr.append(tcMar)


def _compact_paragraph_spacing(paragraph):
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._element.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")


def _set_cell_text_wrap(cell, wrap: bool = True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_no_wrap = tcPr.find(qn("w:noWrap"))
    if existing_no_wrap is not None:
        tcPr.remove(existing_no_wrap)
    if not wrap:
        tcPr.append(OxmlElement("w:noWrap"))


def _configure_cell_paragraph(paragraph, alignment, line_twips: int = 264):
    paragraph.alignment = alignment
    _compact_paragraph_spacing(paragraph)

    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._element.insert(0, pPr)

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), "auto")


def _apply_paragraph_formatting(paragraph, diff: ParagraphDiff):
    info = diff.changed_info or diff.original_info
    if not info:
        return

    pf = paragraph.paragraph_format

    if info.alignment:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        pf.alignment = align_map.get(info.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    if info.space_before_pt is not None:
        pf.space_before = Pt(info.space_before_pt)
    if info.space_after_pt is not None:
        pf.space_after = Pt(info.space_after_pt)
    if info.first_line_indent_pt is not None:
        pf.first_line_indent = Pt(info.first_line_indent_pt)
    if info.left_indent_pt is not None:
        pf.left_indent = Pt(info.left_indent_pt)


def _capture_base_run_formatting(paragraph):
    if not paragraph.runs:
        return {}
    r = paragraph.runs[0]
    return {
        "bold": r.bold,
        "italic": r.italic,
        "font_name": r.font.name,
        "font_size_pt": r.font.size.pt if r.font.size else None,
    }


def _render_delete_in_place(paragraph):
    for run in paragraph.runs:
        run.font.strike = True
        run.font.color.rgb = RED


def _render_delete_in_place_track(paragraph, counter: list[int]) -> None:
    base_rpr = _first_run_rpr_clone(paragraph)
    text = paragraph.text
    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag == qn("w:pPr"):
            continue
        p_elem.remove(child)
    if text:
        _append_del_run(p_elem, text, base_rpr, counter)


def _render_modify_in_place(paragraph, diff: ParagraphDiff):
    base = _capture_base_run_formatting(paragraph)

    p_elem = paragraph._element
    pPr = p_elem.find(qn("w:pPr"))

    for child in list(p_elem):
        if child.tag == qn("w:pPr"):
            continue
        p_elem.remove(child)

    for seg in diff.segments:
        if seg.type == "equal":
            run = paragraph.add_run(seg.text)
        elif seg.type == "delete":
            run = paragraph.add_run(seg.text)
            run.font.strike = True
            run.font.color.rgb = RED
        elif seg.type == "insert":
            run = paragraph.add_run(seg.text)
            run.font.underline = True
            run.font.color.rgb = RED

        if base.get("bold") is not None:
            run.bold = base["bold"]
        if base.get("italic") is not None:
            run.italic = base["italic"]
        if base.get("font_name"):
            run.font.name = base["font_name"]
        if base.get("font_size_pt"):
            run.font.size = Pt(base["font_size_pt"])


def _render_modify_in_place_track(
    paragraph, diff: ParagraphDiff, counter: list[int]
) -> None:
    base_rpr = _first_run_rpr_clone(paragraph)

    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag == qn("w:pPr"):
            continue
        p_elem.remove(child)

    for seg in diff.segments:
        if seg.type == "equal":
            if seg.text:
                _append_plain_run(p_elem, seg.text, base_rpr)
        elif seg.type == "delete":
            if seg.text:
                _append_del_run(p_elem, seg.text, base_rpr, counter)
        elif seg.type == "insert":
            if seg.text:
                _append_ins_run(p_elem, seg.text, base_rpr, counter)


def _render_formatting_in_place(paragraph, diff: ParagraphDiff):
    _apply_paragraph_formatting(paragraph, diff)

    if not diff.has_run_formatting_changes:
        return

    char_pos = 0
    changed_ranges = diff.run_formatting_ranges
    for run in paragraph.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)
        overlaps_change = any(
            r.start < run_end and r.end > run_start for r in changed_ranges
        )
        if overlaps_change:
            run.font.highlight_color = WD_COLOR.YELLOW
        char_pos = run_end


def _render_formatting_in_place_track(
    paragraph, diff: ParagraphDiff, counter: list[int]
) -> None:
    base_rpr = _first_run_rpr_clone(paragraph)
    _apply_paragraph_formatting(paragraph, diff)

    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag == qn("w:pPr"):
            continue
        p_elem.remove(child)

    oi = diff.original_info
    ci = diff.changed_info

    def _runs_del(runs: list[RunInfo]) -> None:
        del_el = OxmlElement("w:del")
        for k, v in _revision_attrs(counter).items():
            del_el.set(k, v)
        for ri in runs:
            if not ri.text:
                continue
            r = OxmlElement("w:r")
            rpr = _rpr_from_run_info(ri)
            if len(rpr):
                r.append(rpr)
            _append_t_to_run(r, ri.text, del_text=True)
            del_el.append(r)
        if len(del_el):
            p_elem.append(del_el)

    def _runs_ins(runs: list[RunInfo]) -> None:
        ins_el = OxmlElement("w:ins")
        for k, v in _revision_attrs(counter).items():
            ins_el.set(k, v)
        for ri in runs:
            if not ri.text:
                continue
            r = OxmlElement("w:r")
            rpr = _rpr_from_run_info(ri)
            if len(rpr):
                r.append(rpr)
            _append_t_to_run(r, ri.text)
            ins_el.append(r)
        if len(ins_el):
            p_elem.append(ins_el)

    if oi and oi.runs:
        _runs_del(oi.runs)
    elif oi and oi.text:
        _append_del_run(p_elem, oi.text, base_rpr, counter)

    if ci and ci.runs:
        _runs_ins(ci.runs)
    elif ci and ci.text:
        _append_ins_run(p_elem, ci.text, base_rpr, counter)


def _insert_paragraph_after(prev_element, diff: ParagraphDiff, doc: Document):
    new_p = OxmlElement("w:p")

    info = diff.changed_info
    if info:
        pPr = OxmlElement("w:pPr")
        if info.alignment:
            align_map = {
                "left": "left",
                "center": "center",
                "right": "right",
                "justify": "both",
            }
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), align_map.get(info.alignment, "left"))
            pPr.append(jc)
        if info.space_before_pt is not None:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:before"), str(int(info.space_before_pt * 20)))
        if info.space_after_pt is not None:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:after"), str(int(info.space_after_pt * 20)))
        if pPr.findall("*"):
            new_p.insert(0, pPr)

    text = info.text if info else ""
    r_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    r_elem.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r_elem.append(t)

    new_p.append(r_elem)

    if prev_element is not None:
        prev_element.addnext(new_p)
    else:
        body = doc.element.body
        body.insert(0, new_p)

    return new_p


def _insert_paragraph_after_track(
    prev_element, diff: ParagraphDiff, doc: Document, counter: list[int]
):
    new_p = OxmlElement("w:p")

    info = diff.changed_info
    if info:
        pPr = OxmlElement("w:pPr")
        if info.alignment:
            align_map = {
                "left": "left",
                "center": "center",
                "right": "right",
                "justify": "both",
            }
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), align_map.get(info.alignment, "left"))
            pPr.append(jc)
        if info.space_before_pt is not None:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:before"), str(int(info.space_before_pt * 20)))
        if info.space_after_pt is not None:
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            spacing.set(qn("w:after"), str(int(info.space_after_pt * 20)))
        if pPr.findall("*"):
            new_p.insert(0, pPr)

    ins = OxmlElement("w:ins")
    for k, v in _revision_attrs(counter).items():
        ins.set(k, v)

    if info and info.runs:
        for ri in info.runs:
            if not ri.text:
                continue
            r_elem = OxmlElement("w:r")
            rpr = _rpr_from_run_info(ri)
            if len(rpr):
                r_elem.append(rpr)
            _append_t_to_run(r_elem, ri.text)
            ins.append(r_elem)
    else:
        text = info.text if info else ""
        r_elem = OxmlElement("w:r")
        _append_t_to_run(r_elem, text)
        ins.append(r_elem)

    new_p.append(ins)

    if prev_element is not None:
        prev_element.addnext(new_p)
    else:
        body = doc.element.body
        body.insert(0, new_p)

    return new_p


def _insert_minimal_legend(doc: Document):
    body = doc.element.body
    first_p = body.find(qn("w:p"))
    if first_p is None:
        return

    legend_p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "120")
    pPr.append(spacing)
    legend_p.append(pPr)

    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    italic1 = OxmlElement("w:i")
    rPr1.append(italic1)
    sz1 = OxmlElement("w:sz")
    sz1.set(qn("w:val"), "16")
    rPr1.append(sz1)
    color1 = OxmlElement("w:color")
    color1.set(qn("w:val"), "999999")
    rPr1.append(color1)
    r1.append(rPr1)
    t1 = OxmlElement("w:t")
    t1.text = "Redline Legend: "
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    legend_p.append(r1)

    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    strike2 = OxmlElement("w:strike")
    rPr2.append(strike2)
    color2 = OxmlElement("w:color")
    color2.set(qn("w:val"), "FF0000")
    rPr2.append(color2)
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "16")
    rPr2.append(sz2)
    r2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = "Strikethrough"
    t2.set(qn("xml:space"), "preserve")
    r2.append(t2)
    legend_p.append(r2)

    r3 = OxmlElement("w:r")
    rPr3 = OxmlElement("w:rPr")
    sz3 = OxmlElement("w:sz")
    sz3.set(qn("w:val"), "16")
    rPr3.append(sz3)
    color3 = OxmlElement("w:color")
    color3.set(qn("w:val"), "999999")
    rPr3.append(color3)
    r3.append(rPr3)
    t3 = OxmlElement("w:t")
    t3.text = " = Deleted  |  "
    t3.set(qn("xml:space"), "preserve")
    r3.append(t3)
    legend_p.append(r3)

    r4 = OxmlElement("w:r")
    rPr4 = OxmlElement("w:rPr")
    u4 = OxmlElement("w:u")
    u4.set(qn("w:val"), "single")
    rPr4.append(u4)
    color4 = OxmlElement("w:color")
    color4.set(qn("w:val"), "FF0000")
    rPr4.append(color4)
    sz4 = OxmlElement("w:sz")
    sz4.set(qn("w:val"), "16")
    rPr4.append(sz4)
    r4.append(rPr4)
    t4 = OxmlElement("w:t")
    t4.text = "Underline"
    t4.set(qn("xml:space"), "preserve")
    r4.append(t4)
    legend_p.append(r4)

    r5 = OxmlElement("w:r")
    rPr5 = OxmlElement("w:rPr")
    sz5 = OxmlElement("w:sz")
    sz5.set(qn("w:val"), "16")
    rPr5.append(sz5)
    color5 = OxmlElement("w:color")
    color5.set(qn("w:val"), "999999")
    rPr5.append(color5)
    r5.append(rPr5)
    t5 = OxmlElement("w:t")
    t5.text = " = Inserted  |  "
    t5.set(qn("xml:space"), "preserve")
    r5.append(t5)
    legend_p.append(r5)

    r6 = OxmlElement("w:r")
    rPr6 = OxmlElement("w:rPr")
    highlight6 = OxmlElement("w:highlight")
    highlight6.set(qn("w:val"), "yellow")
    rPr6.append(highlight6)
    sz6 = OxmlElement("w:sz")
    sz6.set(qn("w:val"), "16")
    rPr6.append(sz6)
    r6.append(rPr6)
    t6 = OxmlElement("w:t")
    t6.text = "Yellow highlight"
    t6.set(qn("xml:space"), "preserve")
    r6.append(t6)
    legend_p.append(r6)

    r7 = OxmlElement("w:r")
    rPr7 = OxmlElement("w:rPr")
    sz7 = OxmlElement("w:sz")
    sz7.set(qn("w:val"), "16")
    rPr7.append(sz7)
    color7 = OxmlElement("w:color")
    color7.set(qn("w:val"), "999999")
    rPr7.append(color7)
    r7.append(rPr7)
    t7 = OxmlElement("w:t")
    t7.text = " = Formatting change"
    t7.set(qn("xml:space"), "preserve")
    r7.append(t7)
    legend_p.append(r7)

    first_p.addprevious(legend_p)


def _start_landscape_section(doc: Document):
    body = doc.element.body
    current_sectPr = body.find(qn("w:sectPr"))

    last_p = None
    for child in list(body):
        if child.tag == qn("w:p"):
            last_p = child
    if last_p is None:
        return

    pPr = last_p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        last_p.insert(0, pPr)

    existing_sect_in_pPr = pPr.find(qn("w:sectPr"))
    if existing_sect_in_pPr is not None:
        pPr.remove(existing_sect_in_pPr)

    if current_sectPr is not None:
        body.remove(current_sectPr)
        pPr.append(current_sectPr)
    else:
        portrait_sectPr = OxmlElement("w:sectPr")
        pgSz = OxmlElement("w:pgSz")
        pgSz.set(qn("w:w"), "12240")
        pgSz.set(qn("w:h"), "15840")
        portrait_sectPr.append(pgSz)
        pgMar = OxmlElement("w:pgMar")
        pgMar.set(qn("w:top"), "1440")
        pgMar.set(qn("w:right"), "1440")
        pgMar.set(qn("w:bottom"), "1440")
        pgMar.set(qn("w:left"), "1440")
        pgMar.set(qn("w:header"), "720")
        pgMar.set(qn("w:footer"), "720")
        pgMar.set(qn("w:gutter"), "0")
        portrait_sectPr.append(pgMar)
        pPr.append(portrait_sectPr)


def _finalize_landscape_section(doc: Document):
    body = doc.element.body
    existing_sectPr = body.find(qn("w:sectPr"))
    if existing_sectPr is not None:
        body.remove(existing_sectPr)

    landscape_sectPr = OxmlElement("w:sectPr")
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "15840")
    pgSz.set(qn("w:h"), "12240")
    pgSz.set(qn("w:orient"), "landscape")
    landscape_sectPr.append(pgSz)
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:top"), "1080")
    pgMar.set(qn("w:right"), "1080")
    pgMar.set(qn("w:bottom"), "1080")
    pgMar.set(qn("w:left"), "1080")
    pgMar.set(qn("w:header"), "720")
    pgMar.set(qn("w:footer"), "720")
    pgMar.set(qn("w:gutter"), "0")
    landscape_sectPr.append(pgMar)
    body.append(landscape_sectPr)


def _render_change_report(doc: Document, changes: list[Change]):
    _start_landscape_section(doc)

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_formatted_run(
        p, "EXHIBIT A \u2014 CHANGE REPORT", bold=True, font_size_pt=16, color=DARK_GRAY
    )

    doc.add_paragraph()

    insertions = sum(1 for c in changes if c.type == ChangeType.INSERTION)
    deletions = sum(1 for c in changes if c.type == ChangeType.DELETION)
    modifications = sum(1 for c in changes if c.type == ChangeType.MODIFICATION)
    formatting = sum(1 for c in changes if c.type == ChangeType.FORMATTING)
    total = len(changes)

    summary_p = doc.add_paragraph()
    _add_formatted_run(
        summary_p,
        f"{total} total changes",
        bold=True,
        font_name=TABLE_FONT,
        font_size_pt=11,
        color=DARK_GRAY,
    )
    _add_formatted_run(
        summary_p,
        (
            f"  |  {insertions} insertions"
            f"  |  {deletions} deletions"
            f"  |  {modifications} modifications"
            f"  |  {formatting} formatting changes"
        ),
        font_name=TABLE_FONT,
        font_size_pt=10,
        color=DARK_GRAY,
    )

    intro_p = doc.add_paragraph()
    _add_formatted_run(
        intro_p,
        "Each change is shown below with the original text beside the revised text.",
        italic=True,
        font_name=TABLE_FONT,
        font_size_pt=9.5,
        color=DARK_GRAY,
    )

    doc.add_paragraph()

    if not changes:
        no_changes_p = doc.add_paragraph()
        no_changes_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_formatted_run(
            no_changes_p,
            "No changes detected between the two documents.",
            italic=True,
            color=DARK_GRAY,
        )
        _finalize_landscape_section(doc)
        return

    for change in changes:
        _render_change_card(doc, change)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    _finalize_landscape_section(doc)


def _truncate_for_cell(text: str, max_len: int = 500) -> str:
    if not text:
        return "\u2014"
    cleaned = " ".join(text.split())
    if len(cleaned) > max_len:
        return cleaned[: max_len - 3] + "..."
    return cleaned


def _format_location(text: str) -> str:
    if text.startswith("Paragraph "):
        return "\u00b6" + text.removeprefix("Paragraph ").strip()
    return text


def _format_detail_text(text: str) -> str:
    if text.startswith('Replaced "') and '" with "' in text:
        old, new = text[len('Replaced "') :].split('" with "', 1)
        new = new.removesuffix('"')
        return f"\"{old}\" -> \"{new}\""

    if text.startswith('Inserted "'):
        inserted = text[len('Inserted "') :].removesuffix('"')
        return f"+ \"{inserted}\""

    if text.startswith('Deleted "'):
        deleted = text[len('Deleted "') :].removesuffix('"')
        return f"- \"{deleted}\""

    return text.replace("; ", ";\n")


def _merge_changed_and_detail(change: Change) -> str:
    changed = _truncate_for_cell(change.new_text, 500)
    detail = _truncate_for_cell(change.formatting_detail, 500)

    if detail == "\u2014":
        return changed

    if changed == "\u2014":
        return detail

    formatted_detail = _format_detail_text(detail)
    return f"{changed}\n\n[{formatted_detail}]"


def _change_type_fill(change_type: ChangeType) -> str:
    return {
        ChangeType.INSERTION: "D5F5E3",
        ChangeType.DELETION: "FADBD8",
        ChangeType.MODIFICATION: "FEF9E7",
        ChangeType.FORMATTING: "D6EAF8",
    }.get(change_type, "FFFFFF")


def _change_type_label(change_type: ChangeType) -> str:
    return {
        ChangeType.INSERTION: "Insertion",
        ChangeType.DELETION: "Deletion",
        ChangeType.MODIFICATION: "Modification",
        ChangeType.FORMATTING: "Formatting Change",
    }.get(change_type, change_type.value.title())


def _clear_cell_paragraphs(cell):
    for paragraph in list(cell.paragraphs):
        p_elem = paragraph._element
        p_elem.getparent().remove(p_elem)


def _add_cell_text(
    cell,
    text: str,
    *,
    font_size_pt: float = 8.5,
    bold: bool = False,
    color: RGBColor | None = None,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
):
    _clear_cell_paragraphs(cell)
    lines = text.splitlines() or [""]
    for idx, line in enumerate(lines):
        paragraph = cell.add_paragraph()
        _configure_cell_paragraph(paragraph, alignment, line_twips=264)
        _add_formatted_run(
            paragraph,
            line or "\u200b",
            bold=bold,
            color=color,
            font_name=TABLE_FONT,
            font_size_pt=font_size_pt,
        )


def _prepare_report_body(text: str) -> str:
    if text == "\u2014":
        return text
    return text.replace("; ", ";\n")


def _render_change_card(doc: Document, change: Change):
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    _set_table_fixed_layout(table)

    col_widths = [4.8, 4.8]
    for row in table.rows:
        for cell, width in zip(row.cells, col_widths):
            cell.width = Inches(width)

    header_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(header_cell, _change_type_fill(change.type))
    _set_cell_border(header_cell)
    _set_cell_margins(header_cell, top=70, bottom=70, left=90, right=90)
    _set_cell_text_wrap(header_cell, wrap=True)
    header_p = header_cell.paragraphs[0]
    _configure_cell_paragraph(header_p, WD_ALIGN_PARAGRAPH.LEFT, line_twips=280)
    _add_formatted_run(
        header_p,
        _change_type_label(change.type),
        bold=True,
        color=DARK_GRAY,
        font_name=TABLE_FONT,
        font_size_pt=10,
    )
    _add_formatted_run(
        header_p,
        f"  |  {_format_location(change.location_desc)}",
        color=DARK_GRAY,
        font_name=TABLE_FONT,
        font_size_pt=10,
    )

    label_row = table.rows[1]
    for label, cell in zip(("Original", "Changed"), label_row.cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, "EAF0F6")
        _set_cell_border(cell)
        _set_cell_margins(cell, top=45, bottom=45, left=80, right=80)
        _set_cell_text_wrap(cell, wrap=True)
        _add_cell_text(
            cell,
            label,
            font_size_pt=8,
            bold=True,
            color=DARK_GRAY,
        )

    body_row = table.rows[2]
    for idx, cell in enumerate(body_row.cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_shading(cell, "FCFBF7")
        _set_cell_border(cell)
        _set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
        _set_cell_text_wrap(cell, wrap=True)
        body_text = (
            _prepare_report_body(_truncate_for_cell(change.original_text, 500))
            if idx == 0
            else _prepare_report_body(_truncate_for_cell(change.new_text, 500))
        )
        _add_cell_text(cell, body_text, font_size_pt=8.5)

    detail = _truncate_for_cell(change.formatting_detail, 500)
    if detail != "\u2014":
        note_row = table.add_row()
        note_cell = note_row.cells[0].merge(note_row.cells[1])
        note_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_shading(note_cell, "F5F7FA")
        _set_cell_border(note_cell)
        _set_cell_margins(note_cell, top=55, bottom=55, left=90, right=90)
        _set_cell_text_wrap(note_cell, wrap=True)
        note_p = note_cell.paragraphs[0]
        _configure_cell_paragraph(note_p, WD_ALIGN_PARAGRAPH.LEFT, line_twips=260)
        _add_formatted_run(
            note_p,
            "Note: ",
            bold=True,
            color=DARK_GRAY,
            font_name=TABLE_FONT,
            font_size_pt=8,
        )
        _add_formatted_run(
            note_p,
            _format_detail_text(detail),
            color=DARK_GRAY,
            font_name=TABLE_FONT,
            font_size_pt=8,
        )


def generate_redline(
    original_path: str,
    changed_path: str,
    output_path: str,
    *,
    output_mode: str = "styled",
) -> None:
    from docx_redline.differ import compare_documents

    if output_mode not in ("styled", "track_changes"):
        raise ValueError(
            f"output_mode must be 'styled' or 'track_changes', got {output_mode!r}"
        )

    doc = Document(original_path)

    orig_paras = [p for p, _ in list_redline_paragraphs(doc)]

    if output_mode == "track_changes":
        _enable_track_revisions(doc)
    else:
        _insert_minimal_legend(doc)

    diffs, changes = compare_documents(original_path, changed_path)

    last_element = None
    rev_counter = [0]

    for diff in diffs:
        if diff.type == "equal":
            oi = diff.original_para_index
            if oi is not None and oi < len(orig_paras):
                last_element = orig_paras[oi]._element

        elif diff.type == "delete":
            oi = diff.original_para_index
            if oi is not None and oi < len(orig_paras):
                para = orig_paras[oi]
                if output_mode == "track_changes":
                    _render_delete_in_place_track(para, rev_counter)
                else:
                    _render_delete_in_place(para)
                last_element = para._element

        elif diff.type == "modify":
            oi = diff.original_para_index
            if oi is not None and oi < len(orig_paras):
                para = orig_paras[oi]
                if output_mode == "track_changes":
                    _render_modify_in_place_track(para, diff, rev_counter)
                else:
                    _render_modify_in_place(para, diff)
                last_element = para._element

        elif diff.type == "formatting":
            oi = diff.original_para_index
            if oi is not None and oi < len(orig_paras):
                para = orig_paras[oi]
                if output_mode == "track_changes":
                    _render_formatting_in_place_track(para, diff, rev_counter)
                else:
                    _render_formatting_in_place(para, diff)
                last_element = para._element

        elif diff.type == "insert":
            if output_mode == "track_changes":
                new_elem = _insert_paragraph_after_track(
                    last_element, diff, doc, rev_counter
                )
            else:
                new_elem = _insert_paragraph_after(last_element, diff, doc)
            last_element = new_elem

    _render_change_report(doc, changes)

    doc.save(output_path)
